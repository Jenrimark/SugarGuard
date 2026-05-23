# -*- coding: utf-8 -*-
"""
SugarGuard AI —— Git 分支策略说明文档生成脚本
输出：GIT_BRANCH_STRATEGY.docx
"""
from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUT_DIR = r"E:\code\Anroid\MyApplication"
os.makedirs(OUT_DIR, exist_ok=True)


# ==================== docx 构造辅助 ====================
doc = Document()

_normal = doc.styles["Normal"]
_normal.font.name = "宋体"
_normal.font.size = Pt(11)
_normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
_normal.paragraph_format.line_spacing = 1.5

for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)


def set_run_font(run, name="宋体", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    sizes = {0: 22, 1: 16, 2: 14, 3: 12, 4: 11}
    set_run_font(run, name="黑体", size=sizes.get(level, 12), bold=True,
                 color=(0x1F, 0x3A, 0x5F))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(text, indent=True, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_bullet(text, size=11):
    p = doc.add_paragraph(style=None)
    run = p.add_run("• " + text)
    set_run_font(run, size=size)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def set_cell_shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "888888")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def make_table(headers, rows, col_widths=None, head_color="1F3A5F",
               zebra=True, font_size=10, align_cols=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, name="黑体", size=font_size, bold=True,
                     color=(0xFF, 0xFF, 0xFF))
        set_cell_shade(c, head_color)
        set_cell_borders(c)
    for r, row in enumerate(rows):
        for j, v in enumerate(row):
            c = table.rows[r + 1].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            if align_cols and j in align_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(v))
            set_run_font(run, size=font_size)
            if zebra and r % 2 == 1:
                set_cell_shade(c, "F5F7FA")
            set_cell_borders(c)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


# ==================== 文档正文 ====================

# 封面
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(100)
r = cover.add_run("SugarGuard AI · 糖知")
set_run_font(r, "黑体", 24, bold=True, color=(0x1F, 0x3A, 0x5F))

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover2.add_run("Git 分支策略说明")
set_run_font(r, "黑体", 28, bold=True, color=(0x26, 0xA6, 0x9A))

for _ in range(3):
    doc.add_paragraph()

info_rows = [
    ("项目名称", "SugarGuard AI（糖知）"),
    ("代码仓库", "https://github.com/xingranju/SugarGuard"),
    ("分支模型", "Git Flow（main / dev / feature / bugfix / hotfix）"),
    ("编制人", "鞠星冉"),
    ("编制日期", "2026-04-18"),
    ("文档版本", "V1.0"),
]
make_table(["项目信息", "内容"], info_rows,
           col_widths=[4, 12], head_color="1F3A5F",
           zebra=True, align_cols=[0])

page_break()

# ===== 一、分支总览 =====
add_heading("一、分支总览", level=1)

add_para(
    "本项目采用 Git Flow 分支模型，定义 main、dev、feature、bugfix、hotfix 五种分支类型，"
    "以保证代码版本的清晰、可追溯和团队协作的有序性。下表为各分支的基本信息。"
)

overview_rows = [
    ["main", "main", "永久", "—", "—"],
    ["dev", "dev", "永久", "main", "main"],
    ["feature", "feature/<模块>-<描述>", "临时", "dev", "dev"],
    ["bugfix", "bugfix/<issue号>-<描述>", "临时", "dev", "dev"],
    ["hotfix", "hotfix/<版本号>-<描述>", "临时", "main", "main + dev"],
]
make_table(["分支类型", "命名规范", "生命周期", "来源分支", "合并目标"],
           overview_rows, col_widths=[2, 5, 2, 3, 4],
           align_cols=[0, 2, 3, 4])

page_break()

# ===== 二、各分支用途详解 =====
add_heading("二、各分支用途详解", level=1)

add_heading("2.1 main（主分支）", level=2)
add_para(
    "main 分支始终保持可发布的稳定代码，每个 commit 对应一个正式发布版本。"
    "禁止直接推送代码到 main，只接受来自 dev（发版合并）或 hotfix（紧急修复）的 Pull Request。"
    "每次合并到 main 后，需要打上版本 Tag（如 v1.0.0）。"
)
add_bullet("保护规则：开启分支保护，要求 PR 合并、至少 1 人 Code Review、CI 全绿。")
add_bullet("版本标签：合并后立即打 Tag，格式 v<主版本>.<次版本>.<修订号>。")
add_bullet("禁止操作：禁止 force push、禁止删除分支、禁止直接 commit。")

add_heading("2.2 dev（开发分支）", level=2)
add_para(
    "dev 分支是日常开发的集成分支，汇聚所有已完成的功能开发和 Bug 修复。"
    "dev 分支应始终保持可编译、核心功能可用的状态。"
    "当 dev 分支达到发布标准（CI 全绿、评审通过）时，向 main 发起 Release PR。"
)
add_bullet("保护规则：建议通过 PR 合并，至少 1 人 Code Review。")
add_bullet("集成节奏：每个 Sprint 结束时，dev 合入 main 作为该 Sprint 的正式发版。")

add_heading("2.3 feature（功能分支）", level=2)
add_para(
    "feature 分支用于开发新功能或较大的需求改进。每个功能分支对应一个或一组关联的用户故事。"
    "功能分支从 dev 创建，开发完成后通过 PR 合入 dev，合入后删除该分支。"
)
add_para("命名示例：", bold=True)
add_bullet("feature/android-drink-record — Android 端饮品记录功能")
add_bullet("feature/backend-ai-chat — 后端 AI 对话接口")
add_bullet("feature/ai-rag-knowledge — AI 服务 RAG 知识库")
add_bullet("feature/android-diary-calendar — 饮食日记日历视图")

add_heading("2.4 bugfix（缺陷修复分支）", level=2)
add_para(
    "bugfix 分支用于修复 dev 分支上发现的非紧急 Bug，即 Sprint 内的常规缺陷修复。"
    "从 dev 创建，修复完成后通过 PR 合入 dev，合入后删除该分支。"
    "分支名应包含 Issue/缺陷编号，以便追溯。"
)
add_para("命名示例：", bold=True)
add_bullet("bugfix/42-login-crash — 修复 #42 登录崩溃")
add_bullet("bugfix/55-meal-time-parse — 修复 #55 用餐时间解析错误")
add_bullet("bugfix/TAPD-1001-jwt-expired — 修复 TAPD #1001 JWT 过期问题")

add_heading("2.5 hotfix（热修复分支）", level=2)
add_para(
    "hotfix 分支用于修复已发布版本（main 分支）上的紧急线上问题，不等待常规发版周期。"
    "从 main 创建，修复完成后同时 PR 合入 main 和 dev，合入 main 后打新的修订版本 Tag，"
    "然后删除该分支。"
)
add_para("命名示例：", bold=True)
add_bullet("hotfix/1.2.1-token-expired — 紧急修复 v1.2 Token 过期问题")
add_bullet("hotfix/1.3.1-db-migration — 紧急修复 v1.3 数据库迁移失败")

page_break()

# ===== 三、合并规则 =====
add_heading("三、合并规则", level=1)

add_heading("3.1 合并方式", level=2)
merge_rows = [
    ["feature → dev", "Squash Merge", "将功能分支的多次提交压缩为一个清晰的 commit，保持 dev 历史简洁"],
    ["bugfix → dev", "Squash Merge", "同上，保持 dev 历史简洁"],
    ["dev → main", "Merge Commit (--no-ff)", "保留完整的发版合并节点，便于追溯每次发版"],
    ["hotfix → main", "Merge Commit (--no-ff)", "保留热修复记录，便于审计"],
    ["hotfix → dev", "Cherry-pick 或 Merge", "确保热修复同步到开发线，避免修复丢失"],
]
make_table(["合并场景", "推荐方式", "说明"],
           merge_rows, col_widths=[4, 4, 8], align_cols=[0, 1])

add_heading("3.2 合并前检查清单", level=2)
add_bullet("CI/CD Pipeline 全部通过（编译、单元测试、Lint 检查）。")
add_bullet("至少 1 人完成 Code Review 并 Approve。")
add_bullet("无合并冲突（若有冲突需在 feature/bugfix 分支上先解决）。")
add_bullet("PR 描述包含关联的 Issue/需求编号（如 TAPD 编号或 GitHub Issue）。")
add_bullet("合入 main 前需在 dev 上完成回归测试，确保无破坏性变更。")

add_heading("3.3 分支保护配置建议", level=2)
protect_rows = [
    ["main", "是", "1+", "build, test", "禁止", "禁止"],
    ["dev", "是", "1+", "build", "禁止", "允许"],
    ["feature/*", "否", "—", "—", "允许", "允许"],
    ["bugfix/*", "否", "—", "—", "允许", "允许"],
    ["hotfix/*", "否", "—", "—", "允许", "允许"],
]
make_table(["分支", "要求 PR", "最少审批人", "必须通过的状态检查", "Force Push", "删除"],
           protect_rows, col_widths=[2.5, 2, 2, 4, 2.5, 2],
           align_cols=[0, 1, 2, 4, 5])

page_break()

# ===== 四、Commit Message 格式 =====
add_heading("四、Commit Message 格式", level=1)

add_para(
    "本项目采用 Conventional Commits 规范，格式统一，便于自动生成 Changelog、语义化版本管理。"
    "所有成员的每次提交必须严格遵循以下格式。"
)

add_heading("4.1 格式规范", level=2)
add_para("<type>(<scope>): <subject>", indent=False, bold=True, size=12)
doc.add_paragraph()
add_para("[可选 body]", indent=False, size=11)
doc.add_paragraph()
add_para("[可选 footer]", indent=False, size=11)

add_heading("4.2 Type 类型说明", level=2)
type_rows = [
    ["feat", "新功能", "feat(android): 添加饮品记录页面"],
    ["fix", "Bug 修复", "fix(backend): 修复 JWT Token 过期未刷新"],
    ["docs", "文档变更", "docs: 更新 API 接口说明"],
    ["style", "代码格式（不影响逻辑）", "style(android): 统一缩进为 4 空格"],
    ["refactor", "重构（非新功能、非修复）", "refactor(ai): 拆分 RAG 查询模块"],
    ["test", "添加/修改测试", "test(backend): 补充用户注册接口测试"],
    ["chore", "构建/工具/依赖变更", "chore: 升级 Spring Boot 至 3.2.0"],
    ["perf", "性能优化", "perf(ai): 优化知识库向量检索速度"],
    ["ci", "CI/CD 配置变更", "ci: 添加 Android APK 构建流水线"],
]
make_table(["Type", "含义", "示例"],
           type_rows, col_widths=[2, 5, 9], align_cols=[0])

add_heading("4.3 Scope 模块标识", level=2)
add_para(
    "Scope 用于标注改动所属模块，为可选字段但推荐填写。"
    "本项目定义以下标准 Scope 值："
)
scope_rows = [
    ["android", "app/", "Android 前端"],
    ["backend", "backend-api/", "Spring Boot 后端"],
    ["ai", "ai-service/", "Python AI 服务"],
    ["scripts", "scripts/", "脚本工具"],
    ["ci", ".github/", "CI/CD 配置"],
    ["docker", "docker-compose.yml 等", "容器编排配置"],
    ["db", "数据库相关", "SQL 脚本、迁移文件"],
]
make_table(["Scope", "对应目录", "说明"],
           scope_rows, col_widths=[3, 5, 8], align_cols=[0])

add_heading("4.4 Subject 要求", level=2)
add_bullet("使用中文或英文均可，团队内统一即可。")
add_bullet("不超过 50 个字符。")
add_bullet("不加句号结尾。")
add_bullet("使用祈使语气（如「添加」而非「添加了」）。")

add_heading("4.5 Body 要求（可选）", level=2)
add_bullet("解释为什么做这个改动，而非做了什么。")
add_bullet("每行不超过 72 个字符。")
add_bullet("与 subject 之间空一行。")

add_heading("4.6 Footer 要求（可选）", level=2)
add_bullet("BREAKING CHANGE: 标注不兼容变更，描述变更内容及迁移方法。")
add_bullet("Closes #issue号: 关联 Issue，合并时自动关闭对应 Issue。")

page_break()

# ===== 五、完整示例 =====
add_heading("五、Commit Message 完整示例", level=1)

add_heading("5.1 新功能提交", level=2)
add_para("feat(android): 添加每日饮品摄入记录功能", indent=False, bold=True)
doc.add_paragraph()
add_para("实现 AddDrinkRecordScreen，支持搜索数据库中的饮品、", indent=False)
add_para("自定义杯量和记录时间，数据同步至后端 API。", indent=False)
doc.add_paragraph()
add_para("Closes #23", indent=False)

add_heading("5.2 Bug 修复提交", level=2)
add_para("fix(backend): 修复分页查询越界导致 500 错误", indent=False, bold=True)
doc.add_paragraph()
add_para("当 page 参数大于实际总页数时，JPA 抛出异常。", indent=False)
add_para("改为返回空列表并正确设置 totalPages。", indent=False)
doc.add_paragraph()
add_para("Closes #42", indent=False)

add_heading("5.3 热修复提交", level=2)
add_para("hotfix(backend): 紧急修复 Token 刷新接口返回 403", indent=False, bold=True)
doc.add_paragraph()
add_para("v1.2.0 上线后发现 SecurityConfig 遗漏了 /api/auth/refresh", indent=False)
add_para("的白名单配置。", indent=False)
doc.add_paragraph()
add_para("Closes #58", indent=False)

page_break()

# ===== 六、典型工作流 =====
add_heading("六、典型工作流", level=1)

add_heading("6.1 开发新功能", level=2)
flow1 = [
    ["1", "切换到 dev 并拉取最新代码", "git checkout dev && git pull origin dev"],
    ["2", "创建功能分支", "git checkout -b feature/android-drink-record"],
    ["3", "开发并提交（可多次）",
     "git commit -m \"feat(android): 添加饮品搜索和记录 UI\"\n"
     "git commit -m \"feat(android): 集成后端饮品 API\""],
    ["4", "推送并创建 PR", "git push -u origin feature/android-drink-record\n→ 在 GitHub 创建 PR: feature → dev"],
    ["5", "Code Review + CI 通过", "Squash Merge 到 dev"],
    ["6", "删除远程功能分支", "合并后自动或手动删除"],
]
make_table(["步骤", "操作", "命令/说明"],
           flow1, col_widths=[1, 5, 10], align_cols=[0])

add_heading("6.2 修复开发分支 Bug", level=2)
flow2 = [
    ["1", "从 dev 创建修复分支", "git checkout dev && git pull\ngit checkout -b bugfix/42-login-crash"],
    ["2", "修复并提交", "git commit -m \"fix(android): 修复登录页空指针崩溃\""],
    ["3", "推送并创建 PR", "git push -u origin bugfix/42-login-crash\n→ PR: bugfix → dev"],
    ["4", "Review 通过后 Squash Merge", "合并后删除分支"],
]
make_table(["步骤", "操作", "命令/说明"],
           flow2, col_widths=[1, 5, 10], align_cols=[0])

add_heading("6.3 修复线上紧急 Bug（Hotfix）", level=2)
flow3 = [
    ["1", "从 main 创建 hotfix 分支", "git checkout main && git pull\ngit checkout -b hotfix/1.2.1-token-expired"],
    ["2", "修复并提交", "git commit -m \"hotfix(backend): 紧急修复 Token 刷新 403\""],
    ["3", "推送并创建两个 PR",
     "git push -u origin hotfix/1.2.1-token-expired\n"
     "→ PR1: hotfix → main\n"
     "→ PR2: hotfix → dev"],
    ["4", "合并后打 Tag", "git checkout main\ngit tag v1.2.1\ngit push origin v1.2.1"],
    ["5", "删除 hotfix 分支", "合并后自动或手动删除"],
]
make_table(["步骤", "操作", "命令/说明"],
           flow3, col_widths=[1, 5, 10], align_cols=[0])

add_heading("6.4 Sprint 发版流程", level=2)
flow4 = [
    ["1", "dev 功能冻结", "所有 feature/bugfix 分支已合入 dev，CI 全绿"],
    ["2", "创建 Release PR", "dev → main，标题写 Release v1.x.0"],
    ["3", "最终 Review 和回归测试", "至少 2 人 Review；关键路径回归测试通过"],
    ["4", "Merge Commit 合入 main", "git merge --no-ff，保留发版节点"],
    ["5", "打 Tag + Release Note", "git tag v1.x.0 && git push origin v1.x.0\n撰写 Release Note（新功能/修复/已知问题）"],
    ["6", "dev 同步", "main 的 merge commit 自动体现在后续 dev 合并中"],
]
make_table(["步骤", "操作", "说明"],
           flow4, col_widths=[1, 5, 10], align_cols=[0])

page_break()

# ===== 七、分支流程图 =====
add_heading("七、分支流程图", level=1)

add_para(
    "下图展示了 Git Flow 的完整分支拓扑关系：main 为稳定发布线，dev 为日常集成线，"
    "feature/bugfix 从 dev 派生并回合到 dev，hotfix 从 main 派生并同时回合到 main 和 dev。"
)

add_para(
    "main ──●───────────────────●──────────●── (v1.0) ──── (v1.1) ── (v1.2.1)\n"
    "       │                   ↑          ↑↑\n"
    "       │              merge│     merge││hotfix\n"
    "       ↓                   │          │↓\n"
    "dev  ──●──●──●──●──●──●──●──●──●──●──●──●──●──\n"
    "          ↑     ↑        ↑        ↑\n"
    "          │     │        │        │\n"
    "feature/A ●──●──┘        │        │\n"
    "                         │        │\n"
    "feature/B    ●──●──●─────┘        │\n"
    "                                  │\n"
    "bugfix/42       ●──●──────────────┘",
    indent=False, size=10
)

page_break()

# ===== 八、版本号规范 =====
add_heading("八、版本号规范", level=1)

add_para(
    "本项目采用语义化版本（Semantic Versioning），格式为 v<MAJOR>.<MINOR>.<PATCH>："
)

ver_rows = [
    ["MAJOR", "不兼容的 API 变更", "v1.0.0 → v2.0.0"],
    ["MINOR", "新增向后兼容的功能", "v1.0.0 → v1.1.0"],
    ["PATCH", "向后兼容的 Bug 修复（含 hotfix）", "v1.0.0 → v1.0.1"],
]
make_table(["版本位", "变更类型", "示例"],
           ver_rows, col_widths=[3, 7, 6], align_cols=[0])

add_heading("8.1 项目版本规划", level=2)
plan_rows = [
    ["v0.1.0-alpha", "Sprint 1 结束", "三端联调通过，阻断 Bug 清零，首次云端部署"],
    ["v0.2.0-beta", "Sprint 2 结束", "核心功能优化完成，数据库扩充到位"],
    ["v1.0.0", "Sprint 3 结束", "正式发布版本，创新功能 2+ 项完成"],
    ["v1.0.x", "按需", "Hotfix 修复版本"],
]
make_table(["版本号", "时间节点", "里程碑"],
           plan_rows, col_widths=[3, 4, 9], align_cols=[0, 1])

page_break()

# ===== 九、常见操作速查 =====
add_heading("九、常见 Git 操作速查", level=1)

cmd_rows = [
    ["查看当前分支", "git branch"],
    ["查看所有分支（含远程）", "git branch -a"],
    ["切换到已有分支", "git checkout <branch-name>"],
    ["创建并切换新分支", "git checkout -b <new-branch>"],
    ["推送新分支到远程", "git push -u origin <branch-name>"],
    ["拉取远程最新代码", "git pull origin <branch-name>"],
    ["合并分支（保留节点）", "git merge --no-ff <branch-name>"],
    ["删除本地已合并分支", "git branch -d <branch-name>"],
    ["删除远程分支", "git push origin --delete <branch-name>"],
    ["打 Tag", "git tag v1.0.0 && git push origin v1.0.0"],
    ["查看提交日志", "git log --oneline --graph --all"],
    ["暂存当前修改", "git stash && git stash pop"],
    ["解决冲突后继续合并", "手动编辑冲突文件 → git add . → git commit"],
]
make_table(["操作", "命令"],
           cmd_rows, col_widths=[6, 10], align_cols=[])

# ===== 结尾 =====
doc.add_paragraph()
add_para(
    "本文档为 SugarGuard AI 项目的 Git 分支管理规范，所有项目成员须严格遵守。"
    "如有调整需求，由组长在团队会议上讨论决定并更新本文档。",
    bold=True
)

tail = doc.add_paragraph()
tail.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = tail.add_run("编制人：鞠星冉    日期：2026-04-18")
set_run_font(r, size=11, bold=True)

# 保存
out_docx = os.path.join(OUT_DIR, "GIT_BRANCH_STRATEGY.docx")
doc.save(out_docx)
print("OK ->", out_docx)
